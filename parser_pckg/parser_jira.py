from bs4 import BeautifulSoup
from copy import deepcopy
import docx
from docx.shared import Inches, Cm
import datetime
from parser_pckg.table_numbers_month import dict_table_number, month_rus


class JiraReport:

    def __init__(self):
        self.__lst_report = []
        self.__html_data = None
        self.__mydoc = docx.Document()
        self.__name = None
        self.__work_time = None
        self.__table = None
        self.__hours_projects = list()
        self.__list_of_projects = set()

    def get_file_data(self, path):
        with open(path, encoding="utf-8", newline="") as f:
            self.__html_data = f.read()

    def __parse_labels(self, labels):
        str_label = str(labels)
        lst_label = str_label.split(",")
        lst_res = []
        for elem in lst_label:
            new_elem = elem.strip()
            if '.' in new_elem:
                lst_res.append(new_elem[: new_elem.find('.')])
                continue
            lst_res.append(new_elem)
        set_ = set(lst_res)
        return ', '.join(list(set_))

    def __parse_data(self):
        soup = BeautifulSoup(self.__html_data, 'lxml')
        get_all_row = soup.find('table', id="issuetable").find_all('tr')
        count = 0
        count_name = 0
        for elem in get_all_row:
            if count == 0:
                count = 1
                continue
            project = elem.find('td', class_="project").get_text().strip()
            summary = elem.find('td', class_="summary").find('p')
            created = elem.find('td', class_="created").get_text()
            updated = elem.find('td', class_="updated").get_text()
            status = elem.find('td', class_="status").find("span").get_text()
            labels = elem.find('td', class_="labels").get_text()
            if not count_name:
                self.__name = elem.find('td', class_="assignee").get_text()
                if self.__name:
                    count_name += 1
            dict_ = {}
            dict_['Project'] = project
            dict_['Summary'] = self.__parser(summary)
            dict_["Created"] = created
            dict_["Updated"] = updated
            dict_["Status"] = status
            dict_["Labels"] = self.__parse_labels(labels)
            self.__lst_report.append(dict_)

    def __parser(self, elem):
        elem = str(elem)
        pos = elem.find("</span>")
        pos_close = elem.find("</p>")
        return elem[pos + len("</span>"): pos_close].strip()

    def fill_document(self):
        self.__parse_data()

    def read_data(self):
        return deepcopy(self.__lst_report)

    def create_doc(self, path):
        count = 0
        self.__add_time_to_project()
        par_head = self.__mydoc.add_paragraph("Приложение 1. Форма отчета\nУтверждено приказом №    от         ")
        par_head.alignment = 2
        self.__mydoc.add_paragraph("ФИО: " + self.__name.strip())
        self.__mydoc.add_paragraph("Количество часов: " + self.__work_time)
        for elem in self.__lst_report:
            count += 1
            str_2_head = "Задача " + str(count) + " :\n"
            str_ = "Проект: " + elem["Project"] + "\n" + "Задача: " + elem["Summary"] + "\n" \
                  + "Затраченное время: " + elem["Time"] + \
                   "\n" + "Статус: " + elem["Status"] + "\n" + "Список проектов: " + elem["Labels"] + "\n"
            self.__mydoc.add_heading(str_2_head, 3)
            self.__mydoc.add_paragraph(str_)
        self.__mydoc.save(r"{}".format(path))

    def create_table_docx(self, path):
        self.__add_time_to_project()
       # print(self.__name.strip())
        par_head = self.__mydoc.add_paragraph("Приложение 1. Форма отчета")
        par_head.alignment = 2
        par_head_2 = self.__mydoc.add_paragraph("Утверждено приказом №     от         .")
        par_head_2.alignment = 2
        self.__mydoc.add_paragraph("Подразделение: ОРВС")
        table_number = dict_table_number.get(self.__name.strip(), "    ")  #словарь с табелями
        self.__mydoc.add_paragraph("ФИО: " + self.__name.strip() + "\nТабельный номер:{}".format(table_number)
                                   + "\nМесяц: {}".format(month_rus[str(datetime.datetime.now().month)])
                                   + "\nГод: {}".format(datetime.datetime.now().year)
                                   + "\nКоличество отработанных часов: " + self.__work_time)
        self.__calculate_total_hours_to_projects()
        self.__chekc_time()
        self.__table = self.__mydoc.add_table(rows=(len(self.__hours_projects) + 1), cols=4)
        self.__table.style = 'Table Grid'
        self.__table.cell(0, 0).text = '№'
        self.__table.cell(0, 1).text = 'Проект'
        self.__table.cell(0, 2).text = 'Время, ч'
        self.__table.cell(0, 3).text = 'Шифр затрат'
        self.__fill_table_all_hours()
        self.__mydoc.add_paragraph("\nПодробная информация по ссылке: http://jira:8080/issues/?filter=11201")
        self.__mydoc.add_paragraph("Руководитель: Шумский А.П")
        self.__mydoc.add_paragraph("Отчет передан по системе 1С")
        self.__mydoc.save(r"{}".format(path))

    def get_name(self):
        return self.__name

    def set_time(self, str_):
        self.__work_time = str_

    def __add_time_to_project(self):
        count_of_task = len(self.__lst_report)
        hours = int(self.__work_time) // count_of_task
        counter = 0
        for elem in self.__lst_report:
            counter += 1
            if counter == count_of_task and count_of_task * hours < int(self.__work_time):
                hours += int(self.__work_time) - count_of_task * hours
            elem["Time"] = str(hours)

    def __set_columns_inches(self):
        for i in range(len(self.__lst_report)):
            for j in range(4):
                if j == 0:
                    self.__table.cell(i, j).width = Inches(1)
                elif j == 1:
                    self.__table.cell(i, j).width = Inches(1)
                elif j == 2:
                    self.__table.cell(i, j).width = Inches(1)
                elif j == 3:
                    self.__table.cell(i, j).width = Inches(2)

    def __fill_table(self):
        for row, elem in enumerate(self.__lst_report, 1):
            for clmn in range(4):
                if clmn == 0:
                    self.__table.cell(row, clmn).text = str(row)
                elif clmn == 1:
                    self.__table.cell(row, clmn).text = elem["Summary"]
                elif clmn == 2:
                    self.__table.cell(row, clmn).text = elem["Time"]
                elif clmn == 3:
                    self.__table.cell(row, clmn).text = elem["Labels"]

    def __calculate_total_hours_to_projects(self):
     #   lst_ = [elem["Labels"].split(",")[0] for elem in self.__lst_report]
       # lst_ = [elem for elem in lst_ if elem.isdigit()]
        lst_ = list()
        for elem in self.__lst_report:
            lst_labels = elem["Labels"].split(",")
            if lst_labels[0].isdigit():
                lst_.append(lst_labels[0].strip())
            elif len(lst_labels) > 1:
                lst_.append(lst_labels[1].strip())
            elif len(lst_labels) == 1:
                lst_.append(lst_labels[0].strip())

        self.__list_of_projects = set(lst_)
        for project in self.__list_of_projects:
            sum_hours = 0
            for elem in self.__lst_report:
                if project == elem["Labels"].split(",")[0] and elem["Labels"].split(",")[0]:
                    sum_hours += int(elem["Time"])
                elif len(elem["Labels"].split(",")) > 1 and project == elem["Labels"].split(",")[1]:
                    sum_hours += int(elem["Time"])
            self.__hours_projects.append({project: sum_hours})
        sum_proverka = 0
        for hours in self.__hours_projects:
            sum_proverka += hours[str(tuple(hours)[0])]
        if sum_proverka < int(self.__work_time):
            self.__hours_projects[-1][str(tuple(self.__hours_projects[-1])[0])] =\
                self.__hours_projects[-1][str(tuple(self.__hours_projects[-1])[0])] + \
                (int(self.__work_time) - sum_proverka)

    def __fill_table_all_hours(self):
  #      print("List", self.__hours_projects)
        for row, elem in enumerate(self.__hours_projects, 1):
            tuple_ = tuple(elem)
            for clmn in range(3):
                 if clmn == 0:
                    self.__table.cell(row, clmn).text = str(row)
                 elif clmn == 1:
                     self.__table.cell(row, clmn).text = tuple_[0]
                 elif clmn == 2:
                    self.__table.cell(row, clmn).text = str(elem[tuple_[0]])

    def __chekc_time(self):
        print(self.__hours_projects)
        key_ = list()
        new_lst = list()
        for elem in self.__hours_projects:
            key_.append(tuple(elem)[0])
        for pos,elem in enumerate(self.__hours_projects):
            if elem[key_[pos]] == 0:
                continue
            new_lst.append(elem)
        self.__hours_projects.clear()
        self.__hours_projects = deepcopy(new_lst)
